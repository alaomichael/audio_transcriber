# import os
# import math
# import io
# import requests
# import gdown
# from flask import Flask, request, jsonify, render_template, send_file
# import speech_recognition as sr
# from pydub import AudioSegment
# import pyttsx3
# import docx
# import PyPDF2
# import zipfile

# # Function to check file existence and size
# def is_file_up_to_date(destination, file_id):
#     if os.path.exists(destination):
#         # Get the size of the existing file
#         existing_size = os.path.getsize(destination)
#         print("existing_size  ",existing_size)
#         # Get the size of the file on Google Drive
#         drive_url = f'https://drive.google.com/uc?id={file_id}&export=download'
#         response = requests.head(drive_url)
#         drive_file_size = int(response.headers.get('Content-Length', 0))
#         print("drive_file_size  ",drive_file_size)
#         return existing_size == drive_file_size
#     return False

# def download_file_from_google_drive(file_id, destination):
#     url = f'https://drive.google.com/uc?id={file_id}'
#     gdown.download(url, destination, quiet=False)

# def download_ffmpeg():
#     file_id = "1gKl_HiRnh8nKOrfm1FQyhHPoANA8GumK"  # Example file ID
#     destination = "ffmpeg.zip"  # Desired destination filename
#     ffmpeg_binaries_dir = "ffmpeg_binaries"  # Directory to extract the binaries

#     # Check if the file already exists
#     if os.path.exists(destination):
#         actual_size = os.path.getsize(destination)
#         expected_size = 59923823  # Update with the correct size in bytes

#         if actual_size == expected_size:
#             print(f"{destination} exists and is the correct size. No download needed.")
#         else:
#             print(f"{destination} exists but size is incorrect. Downloading...")
#             download_file_from_google_drive(file_id, destination)
#     else:
#         print(f"{destination} does not exist. Downloading...")
#         download_file_from_google_drive(file_id, destination)

#     # Check if the destination directory exists, create if not
#     if not os.path.exists(ffmpeg_binaries_dir):
#         os.makedirs(ffmpeg_binaries_dir)

#     # Check if the ffmpeg_binaries directory already contains the extracted files
#     if any(file.endswith(('.exe', '.dll', '.so')) for file in os.listdir(ffmpeg_binaries_dir)):
#         print(f"FFmpeg binaries already extracted to {ffmpeg_binaries_dir}. No extraction needed.")
#     else:
#         # Extract the downloaded ZIP file
#         with zipfile.ZipFile(destination, 'r') as zip_ref:
#             zip_ref.extractall(ffmpeg_binaries_dir)
#             print(f"Extracted FFmpeg binaries to: {ffmpeg_binaries_dir}")

#     # Optionally, remove the ZIP file after extraction
#     # os.remove(destination)  # Uncomment if you want to delete the ZIP file after extraction

# if __name__ == "__main__":
#     download_ffmpeg()


# # Set the path to the ffmpeg binary
# ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg_binaries", "ffmpeg")
# AudioSegment.converter = ffmpeg_path  # Tell pydub to use the local ffmpeg binary

# app = Flask(__name__)

# # Helper function to handle large audio files and transcription
# def transcribe_large_audio(audio_segment, chunk_length_ms=60000):
#     recognizer = sr.Recognizer()
#     audio_length_ms = len(audio_segment)
#     total_chunks = math.ceil(audio_length_ms / chunk_length_ms)
#     transcriptions = []

#     for i in range(total_chunks):
#         start_time = i * chunk_length_ms
#         end_time = start_time + chunk_length_ms
#         chunk = audio_segment[start_time:end_time]

#         # Convert chunk to wav in-memory (no saving to disk)
#         with io.BytesIO() as wav_buffer:
#             chunk.export(wav_buffer, format="wav")
#             wav_buffer.seek(0)

#             with sr.AudioFile(wav_buffer) as source:
#                 audio_data = recognizer.record(source)
#                 try:
#                     text = recognizer.recognize_google(audio_data)
#                     transcriptions.append(text)
#                 except sr.UnknownValueError:
#                     transcriptions.append(f"Chunk {i + 1} could not be understood.")
#                 except sr.RequestError:
#                     transcriptions.append(f"Chunk {i + 1} had an API error.")

#     return " ".join(transcriptions)

# # Function to format the transcription into readable paragraphs
# def format_transcription(transcription):
#     sentences = transcription.split('. ')
#     paragraph_length = 3
#     paragraphs = [" ".join(sentences[i:i + paragraph_length]) for i in range(0, len(sentences), paragraph_length)]
#     formatted_transcription = "<br><br>".join(paragraphs) + "."
#     return formatted_transcription

# # Helper function to convert text to audio using pyttsx3 (offline)
# def convert_text_to_audio(text, filename, voice_id, speed):
#     engine = pyttsx3.init()
#     voices = engine.getProperty('voices')
#     if voice_id < len(voices):
#         engine.setProperty('voice', voices[voice_id].id)
#     engine.setProperty('rate', speed)
#     engine.save_to_file(text, filename)
#     engine.runAndWait()

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page_num in range(len(reader.pages)):
#         text += reader.pages[page_num].extract_text()
#     return text

# # Function to extract text from docx
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# # Frontend route for transcription page
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Frontend route for text-to-audio conversion page
# @app.route('/text-to-audio')
# def text_to_audio_page():
#     engine = pyttsx3.init()
#     voices = engine.getProperty('voices')
#     return render_template('text_to_audio.html', voices=voices)

# # API route for file upload and transcription
# @app.route('/transcribe', methods=['POST'])
# def transcribe_audio():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('index.html', transcription="No file uploaded.")

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('index.html', transcription="No selected file.")

#     if file and file.filename.endswith('.mp3'):
#         try:
#             audio = AudioSegment.from_mp3(file)
#             transcription = transcribe_large_audio(audio)
#             formatted_transcription = format_transcription(transcription)
#         except Exception as e:
#             formatted_transcription = f"Error during transcription: {str(e)}"

#         if request.headers.get('Content-Type') == 'application/json':
#             return jsonify({"transcription": transcription})
#         return render_template('index.html', transcription=formatted_transcription)

#     return jsonify({"error": "Invalid file format. Only MP3 allowed."}), 400 if request.is_json else render_template('index.html', transcription="Invalid file format. Only MP3 allowed.")

# # API route for converting text to audio
# @app.route('/convert-to-audio', methods=['POST'])
# def convert_to_audio():
#     if 'textfile' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

#     file = request.files['textfile']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

#     if file.filename.endswith('.txt'):
#         text = file.read().decode('utf-8')
#     elif file.filename.endswith('.docx'):
#         text = extract_text_from_docx(file)
#     elif file.filename.endswith('.pdf'):
#         text = extract_text_from_pdf(file)
#     else:
#         return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

#     voice_id = int(request.form.get('voice', 0))
#     speed = int(request.form.get('speed', 150))

#     output_filename = "output_audio.mp3"
#     convert_text_to_audio(text, output_filename, voice_id, speed)

#     if request.headers.get('Content-Type') == 'application/json':
#         return jsonify({"message": "Text converted to audio", "filename": output_filename})
#     return send_file(output_filename, as_attachment=True, download_name=output_filename)

# if __name__ == '__main__':
#     app.run(debug=True)


# import os
# import math
# import io
# import requests
# import gdown
# from flask import Flask, request, jsonify, render_template, send_file
# import speech_recognition as sr
# from pydub import AudioSegment
# from gtts import gTTS  # Import gTTS
# import docx
# import PyPDF2
# import zipfile

# # Function to check file existence and size
# def is_file_up_to_date(destination, file_id):
#     if os.path.exists(destination):
#         existing_size = os.path.getsize(destination)
#         drive_url = f'https://drive.google.com/uc?id={file_id}&export=download'
#         response = requests.head(drive_url)
#         drive_file_size = int(response.headers.get('Content-Length', 0))
#         return existing_size == drive_file_size
#     return False

# def download_file_from_google_drive(file_id, destination):
#     url = f'https://drive.google.com/uc?id={file_id}'
#     gdown.download(url, destination, quiet=False)

# def download_ffmpeg():
#     file_id = "1gKl_HiRnh8nKOrfm1FQyhHPoANA8GumK"  # Example file ID
#     destination = "ffmpeg.zip"  # Desired destination filename
#     ffmpeg_binaries_dir = "ffmpeg_binaries"  # Directory to extract the binaries

#     if os.path.exists(destination):
#         actual_size = os.path.getsize(destination)
#         expected_size = 59923823  # Update with the correct size in bytes

#         if actual_size == expected_size:
#             print(f"{destination} exists and is the correct size. No download needed.")
#         else:
#             print(f"{destination} exists but size is incorrect. Downloading...")
#             download_file_from_google_drive(file_id, destination)
#     else:
#         print(f"{destination} does not exist. Downloading...")
#         download_file_from_google_drive(file_id, destination)

#     if not os.path.exists(ffmpeg_binaries_dir):
#         os.makedirs(ffmpeg_binaries_dir)

#     if any(file.endswith(('.exe', '.dll', '.so')) for file in os.listdir(ffmpeg_binaries_dir)):
#         print(f"FFmpeg binaries already extracted to {ffmpeg_binaries_dir}. No extraction needed.")
#     else:
#         with zipfile.ZipFile(destination, 'r') as zip_ref:
#             zip_ref.extractall(ffmpeg_binaries_dir)
#             print(f"Extracted FFmpeg binaries to: {ffmpeg_binaries_dir}")

# if __name__ == "__main__":
#     download_ffmpeg()

# # Set the path to the ffmpeg binary
# ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg_binaries", "ffmpeg")
# AudioSegment.converter = ffmpeg_path  # Tell pydub to use the local ffmpeg binary

# app = Flask(__name__)

# # Helper function to handle large audio files and transcription
# def transcribe_large_audio(audio_segment, chunk_length_ms=60000):
#     recognizer = sr.Recognizer()
#     audio_length_ms = len(audio_segment)
#     total_chunks = math.ceil(audio_length_ms / chunk_length_ms)
#     transcriptions = []

#     for i in range(total_chunks):
#         start_time = i * chunk_length_ms
#         end_time = start_time + chunk_length_ms
#         chunk = audio_segment[start_time:end_time]

#         with io.BytesIO() as wav_buffer:
#             chunk.export(wav_buffer, format="wav")
#             wav_buffer.seek(0)

#             with sr.AudioFile(wav_buffer) as source:
#                 audio_data = recognizer.record(source)
#                 try:
#                     text = recognizer.recognize_google(audio_data)
#                     transcriptions.append(text)
#                 except sr.UnknownValueError:
#                     transcriptions.append(f"Chunk {i + 1} could not be understood.")
#                 except sr.RequestError:
#                     transcriptions.append(f"Chunk {i + 1} had an API error.")

#     return " ".join(transcriptions)

# # Function to format the transcription into readable paragraphs
# def format_transcription(transcription):
#     sentences = transcription.split('. ')
#     paragraph_length = 3
#     paragraphs = [" ".join(sentences[i:i + paragraph_length]) for i in range(0, len(sentences), paragraph_length)]
#     formatted_transcription = "<br><br>".join(paragraphs) + "."
#     return formatted_transcription

# # Helper function to convert text to audio using gTTS (online)
# def convert_text_to_audio(text, filename, lang='en'):
#     tts = gTTS(text, lang=lang)
#     tts.save(filename)

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page_num in range(len(reader.pages)):
#         text += reader.pages[page_num].extract_text()
#     return text

# # Function to extract text from docx
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# # Frontend route for transcription page
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Frontend route for text-to-audio conversion page
# @app.route('/text-to-audio')
# def text_to_audio_page():
#     return render_template('text_to_audio.html')

# # API route for file upload and transcription
# @app.route('/transcribe', methods=['POST'])
# def transcribe_audio():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('index.html', transcription="No file uploaded.")

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('index.html', transcription="No selected file.")

#     if file and file.filename.endswith('.mp3'):
#         try:
#             audio = AudioSegment.from_mp3(file)
#             transcription = transcribe_large_audio(audio)
#             formatted_transcription = format_transcription(transcription)
#         except Exception as e:
#             formatted_transcription = f"Error during transcription: {str(e)}"

#         if request.headers.get('Content-Type') == 'application/json':
#             return jsonify({"transcription": transcription})
#         return render_template('index.html', transcription=formatted_transcription)

#     return jsonify({"error": "Invalid file format. Only MP3 allowed."}), 400 if request.is_json else render_template('index.html', transcription="Invalid file format. Only MP3 allowed.")

# # API route for converting text to audio
# @app.route('/convert-to-audio', methods=['POST'])
# def convert_to_audio():
#     if 'textfile' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

#     file = request.files['textfile']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

#     if file.filename.endswith('.txt'):
#         text = file.read().decode('utf-8')
#     elif file.filename.endswith('.docx'):
#         text = extract_text_from_docx(file)
#     elif file.filename.endswith('.pdf'):
#         text = extract_text_from_pdf(file)
#     else:
#         return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

#     output_filename = "output_audio.mp3"
#     convert_text_to_audio(text, output_filename)

#     if request.headers.get('Content-Type') == 'application/json':
#         return jsonify({"message": "Text converted to audio", "filename": output_filename})
#     return send_file(output_filename, as_attachment=True, download_name=output_filename)

# if __name__ == '__main__':
#     app.run(debug=True)

# import os
# import math
# import io
# import requests
# import gdown
# from flask import Flask, request, jsonify, render_template, send_file
# import speech_recognition as sr
# from pydub import AudioSegment
# import pyttsx3
# from gtts import gTTS
# import docx
# import PyPDF2
# import zipfile

# # Function to check file existence and size
# def is_file_up_to_date(destination, file_id):
#     if os.path.exists(destination):
#         existing_size = os.path.getsize(destination)
#         drive_url = f'https://drive.google.com/uc?id={file_id}&export=download'
#         response = requests.head(drive_url)
#         drive_file_size = int(response.headers.get('Content-Length', 0))
#         return existing_size == drive_file_size
#     return False

# def download_file_from_google_drive(file_id, destination):
#     url = f'https://drive.google.com/uc?id={file_id}'
#     gdown.download(url, destination, quiet=False)

# def download_ffmpeg():
#     file_id = "1gKl_HiRnh8nKOrfm1FQyhHPoANA8GumK"  # Example file ID
#     destination = "ffmpeg.zip"  # Desired destination filename
#     ffmpeg_binaries_dir = "ffmpeg_binaries"  # Directory to extract the binaries

#     # Check if the file already exists
#     if os.path.exists(destination):
#         actual_size = os.path.getsize(destination)
#         expected_size = 59923823  # Update with the correct size in bytes

#         if actual_size == expected_size:
#             print(f"{destination} exists and is the correct size. No download needed.")
#         else:
#             print(f"{destination} exists but size is incorrect. Downloading...")
#             download_file_from_google_drive(file_id, destination)
#     else:
#         print(f"{destination} does not exist. Downloading...")
#         download_file_from_google_drive(file_id, destination)

#     # Check if the destination directory exists, create if not
#     if not os.path.exists(ffmpeg_binaries_dir):
#         os.makedirs(ffmpeg_binaries_dir)

#     # Check if the ffmpeg_binaries directory already contains the extracted files
#     if any(file.endswith(('.exe', '.dll', '.so')) for file in os.listdir(ffmpeg_binaries_dir)):
#         print(f"FFmpeg binaries already extracted to {ffmpeg_binaries_dir}. No extraction needed.")
#     else:
#         # Extract the downloaded ZIP file
#         with zipfile.ZipFile(destination, 'r') as zip_ref:
#             zip_ref.extractall(ffmpeg_binaries_dir)
#             print(f"Extracted FFmpeg binaries to: {ffmpeg_binaries_dir}")

# if __name__ == "__main__":
#     download_ffmpeg()

# # Set the path to the ffmpeg binary
# ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg_binaries", "ffmpeg")
# AudioSegment.converter = ffmpeg_path  # Tell pydub to use the local ffmpeg binary

# app = Flask(__name__)

# # Helper function to handle large audio files and transcription
# def transcribe_large_audio(audio_segment, chunk_length_ms=60000):
#     recognizer = sr.Recognizer()
#     audio_length_ms = len(audio_segment)
#     total_chunks = math.ceil(audio_length_ms / chunk_length_ms)
#     transcriptions = []

#     for i in range(total_chunks):
#         start_time = i * chunk_length_ms
#         end_time = start_time + chunk_length_ms
#         chunk = audio_segment[start_time:end_time]

#         # Convert chunk to wav in-memory (no saving to disk)
#         with io.BytesIO() as wav_buffer:
#             chunk.export(wav_buffer, format="wav")
#             wav_buffer.seek(0)

#             with sr.AudioFile(wav_buffer) as source:
#                 audio_data = recognizer.record(source)
#                 try:
#                     text = recognizer.recognize_google(audio_data)
#                     transcriptions.append(text)
#                 except sr.UnknownValueError:
#                     transcriptions.append(f"Chunk {i + 1} could not be understood.")
#                 except sr.RequestError:
#                     transcriptions.append(f"Chunk {i + 1} had an API error.")

#     return " ".join(transcriptions)

# # Function to format the transcription into readable paragraphs
# def format_transcription(transcription):
#     sentences = transcription.split('. ')
#     paragraph_length = 3
#     paragraphs = [" ".join(sentences[i:i + paragraph_length]) for i in range(0, len(sentences), paragraph_length)]
#     formatted_transcription = "<br><br>".join(paragraphs) + "."
#     return formatted_transcription

# # Helper function to convert text to audio using pyttsx3 (offline)
# def convert_text_to_audio_with_pyttsx3(text, filename, voice_id, speed):
#     engine = pyttsx3.init()
#     voices = engine.getProperty('voices')
#     if voice_id < len(voices):
#         engine.setProperty('voice', voices[voice_id].id)
#     engine.setProperty('rate', speed)
#     engine.save_to_file(text, filename)
#     engine.runAndWait()

# # Helper function to convert text to audio using gTTS (online)
# def convert_text_to_audio_with_gtts(text, filename):
#     tts = gTTS(text, lang='en')
#     tts.save(filename)

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page_num in range(len(reader.pages)):
#         text += reader.pages[page_num].extract_text()
#     return text

# # Function to extract text from docx
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# # Frontend route for transcription page
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Frontend route for text-to-audio conversion page
# @app.route('/text-to-audio')
# def text_to_audio_page():
#     engine = pyttsx3.init()
#     voices = engine.getProperty('voices')
#     return render_template('text_to_audio.html', voices=voices)

# # API route for file upload and transcription
# @app.route('/transcribe', methods=['POST'])
# def transcribe_audio():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('index.html', transcription="No file uploaded.")

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('index.html', transcription="No selected file.")

#     if file and file.filename.endswith('.mp3'):
#         try:
#             audio = AudioSegment.from_mp3(file)
#             transcription = transcribe_large_audio(audio)
#             formatted_transcription = format_transcription(transcription)
#         except Exception as e:
#             formatted_transcription = f"Error during transcription: {str(e)}"

#         if request.headers.get('Content-Type') == 'application/json':
#             return jsonify({"transcription": transcription})
#         return render_template('index.html', transcription=formatted_transcription)

#     return jsonify({"error": "Invalid file format. Only MP3 allowed."}), 400 if request.is_json else render_template('index.html', transcription="Invalid file format. Only MP3 allowed.")

# # API route for converting text to audio
# @app.route('/convert-to-audio', methods=['POST'])
# def convert_to_audio():
#     if 'textfile' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

#     file = request.files['textfile']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

#     if file.filename.endswith('.txt'):
#         text = file.read().decode('utf-8')
#     elif file.filename.endswith('.docx'):
#         text = extract_text_from_docx(file)
#     elif file.filename.endswith('.pdf'):
#         text = extract_text_from_pdf(file)
#     else:
#         return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

#     # Retrieve voice_id and speed from the form
#     voice_id = int(request.form.get('voice', 0))
#     speed = int(request.form.get('speed', 150))

#     output_filename_pyttsx3 = "output_audio_pyttsx3.mp3"
#     convert_text_to_audio_with_pyttsx3(text, output_filename_pyttsx3, voice_id, speed)

#     output_filename_gtts = "output_audio_gtts.mp3"
#     convert_text_to_audio_with_gtts(text, output_filename_gtts)

#     # Respond with both audio files
#     if request.headers.get('Content-Type') == 'application/json':
#         return jsonify({
#             "message": "Text converted to audio",
#             "pyttsx3_filename": output_filename_pyttsx3,
#             "gtts_filename": output_filename_gtts
#         })
#     return send_file(output_filename_pyttsx3, as_attachment=True, download_name=output_filename_pyttsx3)

# if __name__ == '__main__':
#     app.run(debug=True)



# import os
# import math
# import io
# import requests
# import gdown
# from flask import Flask, request, jsonify, render_template, send_file
# import speech_recognition as sr
# from pydub import AudioSegment
# from gtts import gTTS
# import docx
# import PyPDF2
# import zipfile

# # Function to check file existence and size
# def is_file_up_to_date(destination, file_id):
#     if os.path.exists(destination):
#         existing_size = os.path.getsize(destination)
#         drive_url = f'https://drive.google.com/uc?id={file_id}&export=download'
#         response = requests.head(drive_url)
#         drive_file_size = int(response.headers.get('Content-Length', 0))
#         return existing_size == drive_file_size
#     return False

# def download_file_from_google_drive(file_id, destination):
#     url = f'https://drive.google.com/uc?id={file_id}'
#     gdown.download(url, destination, quiet=False)

# def download_ffmpeg():
#     file_id = "1gKl_HiRnh8nKOrfm1FQyhHPoANA8GumK"  # Example file ID
#     destination = "ffmpeg.zip"  # Desired destination filename
#     ffmpeg_binaries_dir = "ffmpeg_binaries"  # Directory to extract the binaries

#     # Check if the file already exists
#     if os.path.exists(destination):
#         actual_size = os.path.getsize(destination)
#         expected_size = 59923823  # Update with the correct size in bytes

#         if actual_size == expected_size:
#             print(f"{destination} exists and is the correct size. No download needed.")
#         else:
#             print(f"{destination} exists but size is incorrect. Downloading...")
#             download_file_from_google_drive(file_id, destination)
#     else:
#         print(f"{destination} does not exist. Downloading...")
#         download_file_from_google_drive(file_id, destination)

#     # Check if the destination directory exists, create if not
#     if not os.path.exists(ffmpeg_binaries_dir):
#         os.makedirs(ffmpeg_binaries_dir)

#     # Check if the ffmpeg_binaries directory already contains the extracted files
#     if any(file.endswith(('.exe', '.dll', '.so')) for file in os.listdir(ffmpeg_binaries_dir)):
#         print(f"FFmpeg binaries already extracted to {ffmpeg_binaries_dir}. No extraction needed.")
#     else:
#         # Extract the downloaded ZIP file
#         with zipfile.ZipFile(destination, 'r') as zip_ref:
#             zip_ref.extractall(ffmpeg_binaries_dir)
#             print(f"Extracted FFmpeg binaries to: {ffmpeg_binaries_dir}")

# if __name__ == "__main__":
#     download_ffmpeg()

# # Set the path to the ffmpeg binary
# ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg_binaries", "ffmpeg")
# AudioSegment.converter = ffmpeg_path  # Tell pydub to use the local ffmpeg binary

# app = Flask(__name__)

# # Helper function to handle large audio files and transcription
# def transcribe_large_audio(audio_segment, chunk_length_ms=60000):
#     recognizer = sr.Recognizer()
#     audio_length_ms = len(audio_segment)
#     total_chunks = math.ceil(audio_length_ms / chunk_length_ms)
#     transcriptions = []

#     for i in range(total_chunks):
#         start_time = i * chunk_length_ms
#         end_time = start_time + chunk_length_ms
#         chunk = audio_segment[start_time:end_time]

#         # Convert chunk to wav in-memory (no saving to disk)
#         with io.BytesIO() as wav_buffer:
#             chunk.export(wav_buffer, format="wav")
#             wav_buffer.seek(0)

#             with sr.AudioFile(wav_buffer) as source:
#                 audio_data = recognizer.record(source)
#                 try:
#                     text = recognizer.recognize_google(audio_data)
#                     transcriptions.append(text)
#                 except sr.UnknownValueError:
#                     transcriptions.append(f"Chunk {i + 1} could not be understood.")
#                 except sr.RequestError:
#                     transcriptions.append(f"Chunk {i + 1} had an API error.")

#     return " ".join(transcriptions)

# # Function to format the transcription into readable paragraphs
# def format_transcription(transcription):
#     sentences = transcription.split('. ')
#     paragraph_length = 3
#     paragraphs = [" ".join(sentences[i:i + paragraph_length]) for i in range(0, len(sentences), paragraph_length)]
#     formatted_transcription = "<br><br>".join(paragraphs) + "."
#     return formatted_transcription

# # Helper function to convert text to audio using gTTS (online)
# def convert_text_to_audio_with_gtts(text, filename):
#     tts = gTTS(text, lang='en')
#     tts.save(filename)

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     text = ""
#     for page_num in range(len(reader.pages)):
#         text += reader.pages[page_num].extract_text()
#     return text

# # Function to extract text from docx
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# # Frontend route for transcription page
# @app.route('/')
# def index():
#     return render_template('index.html')

# # Frontend route for text-to-audio conversion page
# @app.route('/text-to-audio')
# def text_to_audio_page():
#     return render_template('text_to_audio.html')

# # API route for file upload and transcription
# @app.route('/transcribe', methods=['POST'])
# def transcribe_audio():
#     if 'file' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('index.html', transcription="No file uploaded.")

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('index.html', transcription="No selected file.")

#     if file and file.filename.endswith('.mp3'):
#         try:
#             audio = AudioSegment.from_mp3(file)
#             transcription = transcribe_large_audio(audio)
#             formatted_transcription = format_transcription(transcription)
#         except Exception as e:
#             formatted_transcription = f"Error during transcription: {str(e)}"

#         if request.headers.get('Content-Type') == 'application/json':
#             return jsonify({"transcription": transcription})
#         return render_template('index.html', transcription=formatted_transcription)

#     return jsonify({"error": "Invalid file format. Only MP3 allowed."}), 400 if request.is_json else render_template('index.html', transcription="Invalid file format. Only MP3 allowed.")

# # API route for converting text to audio
# @app.route('/convert-to-audio', methods=['POST'])
# # def convert_to_audio():
# #     if 'textfile' not in request.files:
# #         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

# #     file = request.files['textfile']
# #     if file.filename == '':
# #         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

# #     if file.filename.endswith('.txt'):
# #         text = file.read().decode('utf-8')
# #     elif file.filename.endswith('.docx'):
# #         text = extract_text_from_docx(file)
# #     elif file.filename.endswith('.pdf'):
# #         text = extract_text_from_pdf(file)
# #     else:
# #         return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

# #     output_filename_gtts = "output_audio_gtts.mp3"
# #     convert_text_to_audio_with_gtts(text, output_filename_gtts)

# #     # Respond with the audio file
# #     if request.headers.get('Content-Type') == 'application/json':
# #         return jsonify({
# #             "message": "Text converted to audio",
# #             "gtts_filename": output_filename_gtts
# #         })
# #     return send_file(output_filename_gtts, as_attachment=True, download_name=output_filename_gtts)
# def convert_to_audio():
#     if 'textfile' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

#     file = request.files['textfile']
#     if file.filename == '':
#         return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

#     if file.filename.endswith('.txt'):
#         text = file.read().decode('utf-8')
#     elif file.filename.endswith('.docx'):
#         text = extract_text_from_docx(file)
#     elif file.filename.endswith('.pdf'):
#         text = extract_text_from_pdf(file)
#     else:
#         return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

#     # Create an in-memory bytes buffer
#     audio_stream = io.BytesIO()
    
#     # Convert text to audio using gTTS and save to the in-memory buffer
#     tts = gTTS(text=text, lang='en')
#     tts.save(audio_stream)
#     audio_stream.seek(0)  # Rewind the stream to the beginning

#     # Respond with the audio file
#     if request.headers.get('Content-Type') == 'application/json':
#         return jsonify({
#             "message": "Text converted to audio",
#             "audio_available": True  # Indicating audio is ready for download
#         })

#     # Send the audio stream as a response
#     return send_file(audio_stream, mimetype='audio/mpeg', as_attachment=True, download_name='output_audio.mp3')


# if __name__ == '__main__':
#     app.run(debug=True)

import os
import math
import io
import requests
import gdown
from flask import Flask, request, jsonify, render_template, send_file
import speech_recognition as sr
from pydub import AudioSegment
import pyttsx3
import docx
import PyPDF2
import zipfile

# Function to check file existence and size
def is_file_up_to_date(destination, file_id):
    if os.path.exists(destination):
        existing_size = os.path.getsize(destination)
        drive_url = f'https://drive.google.com/uc?id={file_id}&export=download'
        response = requests.head(drive_url)
        drive_file_size = int(response.headers.get('Content-Length', 0))
        return existing_size == drive_file_size
    return False

def download_file_from_google_drive(file_id, destination):
    url = f'https://drive.google.com/uc?id={file_id}'
    gdown.download(url, destination, quiet=False)

def download_ffmpeg():
    file_id = "1gKl_HiRnh8nKOrfm1FQyhHPoANA8GumK"  # Example file ID
    destination = "ffmpeg.zip"  # Desired destination filename
    ffmpeg_binaries_dir = "ffmpeg_binaries"  # Directory to extract the binaries

    # Check if the file already exists
    if os.path.exists(destination):
        actual_size = os.path.getsize(destination)
        expected_size = 59923823  # Update with the correct size in bytes

        if actual_size == expected_size:
            print(f"{destination} exists and is the correct size. No download needed.")
        else:
            print(f"{destination} exists but size is incorrect. Downloading...")
            download_file_from_google_drive(file_id, destination)
    else:
        print(f"{destination} does not exist. Downloading...")
        download_file_from_google_drive(file_id, destination)

    # Check if the destination directory exists, create if not
    if not os.path.exists(ffmpeg_binaries_dir):
        os.makedirs(ffmpeg_binaries_dir)

    # Check if the ffmpeg_binaries directory already contains the extracted files
    if any(file.endswith(('.exe', '.dll', '.so')) for file in os.listdir(ffmpeg_binaries_dir)):
        print(f"FFmpeg binaries already extracted to {ffmpeg_binaries_dir}. No extraction needed.")
    else:
        # Extract the downloaded ZIP file
        with zipfile.ZipFile(destination, 'r') as zip_ref:
            zip_ref.extractall(ffmpeg_binaries_dir)
            print(f"Extracted FFmpeg binaries to: {ffmpeg_binaries_dir}")

# Download FFmpeg binaries
download_ffmpeg()

# Set the path to the ffmpeg binary
ffmpeg_path = os.path.join(os.getcwd(), "ffmpeg_binaries", "ffmpeg")
AudioSegment.converter = ffmpeg_path  # Tell pydub to use the local ffmpeg binary

app = Flask(__name__)

# Helper function to handle large audio files and transcription
def transcribe_large_audio(audio_segment, chunk_length_ms=60000):
    recognizer = sr.Recognizer()
    audio_length_ms = len(audio_segment)
    total_chunks = math.ceil(audio_length_ms / chunk_length_ms)
    transcriptions = []

    for i in range(total_chunks):
        start_time = i * chunk_length_ms
        end_time = start_time + chunk_length_ms
        chunk = audio_segment[start_time:end_time]

        # Convert chunk to wav in-memory (no saving to disk)
        with io.BytesIO() as wav_buffer:
            chunk.export(wav_buffer, format="wav")
            wav_buffer.seek(0)

            with sr.AudioFile(wav_buffer) as source:
                audio_data = recognizer.record(source)
                try:
                    text = recognizer.recognize_google(audio_data)
                    transcriptions.append(text)
                except sr.UnknownValueError:
                    transcriptions.append(f"Chunk {i + 1} could not be understood.")
                except sr.RequestError:
                    transcriptions.append(f"Chunk {i + 1} had an API error.")

    return " ".join(transcriptions)

# Function to format the transcription into readable paragraphs
def format_transcription(transcription):
    sentences = transcription.split('. ')
    paragraph_length = 3
    paragraphs = [" ".join(sentences[i:i + paragraph_length]) for i in range(0, len(sentences), paragraph_length)]
    formatted_transcription = "<br><br>".join(paragraphs) + "."
    return formatted_transcription

# Helper function to convert text to audio using pyttsx3 (offline)
def convert_text_to_audio_with_pyttsx3(text, filename, voice_id=None, speed=150):
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    
    # If voice_id is None or out of range, default to the first voice
    if voice_id is None or voice_id >= len(voices):
        voice_id = 0  # Default to the first voice
    
    engine.setProperty('voice', voices[voice_id].id)
    engine.setProperty('rate', speed)
    engine.save_to_file(text, filename)
    engine.runAndWait()

# Function to extract text from PDF
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(reader.pages)):
        text += reader.pages[page_num].extract_text()
    return text

# Function to extract text from docx
def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# Frontend route for transcription page
@app.route('/')
def index():
    return render_template('index.html')

# Frontend route for text-to-audio conversion page
@app.route('/text-to-audio')
def text_to_audio_page():
    return render_template('text_to_audio.html')

# API route for file upload and transcription
@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('index.html', transcription="No file uploaded.")

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('index.html', transcription="No selected file.")

    if file and file.filename.endswith('.mp3'):
        try:
            audio = AudioSegment.from_mp3(file)
            transcription = transcribe_large_audio(audio)
            formatted_transcription = format_transcription(transcription)
        except Exception as e:
            formatted_transcription = f"Error during transcription: {str(e)}"

        if request.headers.get('Content-Type') == 'application/json':
            return jsonify({"transcription": transcription})
        return render_template('index.html', transcription=formatted_transcription)

    return jsonify({"error": "Invalid file format. Only MP3 allowed."}), 400 if request.is_json else render_template('index.html', transcription="Invalid file format. Only MP3 allowed.")

# API route for converting text to audio
@app.route('/convert-to-audio', methods=['POST'])
def convert_to_audio():
    if 'textfile' not in request.files:
        return jsonify({"error": "No file uploaded."}), 400 if request.is_json else render_template('text_to_audio.html', error="No file uploaded.")

    file = request.files['textfile']
    if file.filename == '':
        return jsonify({"error": "No selected file."}), 400 if request.is_json else render_template('text_to_audio.html', error="No selected file.")

    if file.filename.endswith('.txt'):
        text = file.read().decode('utf-8')
    elif file.filename.endswith('.docx'):
        text = extract_text_from_docx(file)
    elif file.filename.endswith('.pdf'):
        text = extract_text_from_pdf(file)
    else:
        return jsonify({"error": "Invalid file format. Only txt, docx, and pdf allowed."}), 400 if request.is_json else render_template('text_to_audio.html', error="Invalid file format. Only txt, docx, and pdf allowed.")

    # Create an in-memory bytes buffer
    output_filename = "output_audio.mp3"
    convert_text_to_audio_with_pyttsx3(text, output_filename)

    # Open the audio file to send it as a response
    with open(output_filename, 'rb') as audio_file:
        return send_file(audio_file, mimetype='audio/mpeg', as_attachment=True, download_name='output_audio.mp3')

if __name__ == '__main__':
    app.run(debug=True)
